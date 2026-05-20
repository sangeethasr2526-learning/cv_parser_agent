"""
CV Parser Tools
---------------
7 tools the LLM agent can call.
Each returns a plain dict (gets JSON-serialised before being sent back to the LLM).
"""

import os
import re
from datetime import datetime
from pathlib import Path


# ── Tool 1: detect_file_type ──────────────────────────────────────────────────
def detect_file_type(file_path: str) -> dict:
    """
    Looks at the file extension (and magic bytes for safety)
    and returns which reader tool to use next.
    """
    path = Path(file_path)

    if not path.exists():
        return {"error": f"File not found: {file_path}"}

    ext = path.suffix.lower()

    type_map = {
        ".pdf":  "pdf",
        ".docx": "docx",
        ".doc":  "docx",
        ".txt":  "text",
        ".png":  "image",
        ".jpg":  "image",
        ".jpeg": "image",
        ".tiff": "image",
        ".tif":  "image",
        ".bmp":  "image",
    }

    file_type = type_map.get(ext, "unknown")

    # Basic magic byte check for PDF (starts with %PDF)
    if file_type == "unknown" or ext == "":
        try:
            with open(file_path, "rb") as f:
                header = f.read(4)
            if header == b"%PDF":
                file_type = "pdf"
        except Exception:
            pass

    tool_to_call = {
        "pdf":     "read_pdf",
        "docx":    "read_docx",
        "image":   "ocr_image",
        "text":    "read_text",
        "unknown": "unknown",
    }.get(file_type, "unknown")

    return {
        "file_path":    file_path,
        "file_type":    file_type,
        "tool_to_call": tool_to_call,
        "file_size_kb": round(path.stat().st_size / 1024, 1),
    }


# ── Tool 2: read_pdf ──────────────────────────────────────────────────────────
def read_pdf(file_path: str) -> dict:
    """
    Extracts all text from a PDF using pdfplumber.
    Falls back to pypdf if pdfplumber isn't installed.
    """
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            num_pages = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

        full_text = "\n\n".join(text_parts).strip()

        if not full_text:
            return {
                "success": False,
                "error": "No text extracted — file may be a scanned PDF. Try ocr_image instead.",
                "pages": num_pages,
            }

        return {
            "success":    True,
            "text":       full_text,
            "pages":      num_pages,
            "char_count": len(full_text),
        }

    except ImportError:
        return {"error": "pdfplumber not installed. Run: pip install pdfplumber"}
    except Exception as e:
        return {"error": str(e)}


# ── Tool 3: read_docx ─────────────────────────────────────────────────────────
def read_docx(file_path: str) -> dict:
    """
    Extracts all text from a Word .docx file using python-docx.
    """
    try:
        from docx import Document
        doc = Document(file_path)

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also grab text from tables (some CVs use table layouts)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)

        full_text = "\n".join(paragraphs).strip()

        return {
            "success":    True,
            "text":       full_text,
            "char_count": len(full_text),
        }

    except ImportError:
        return {"error": "python-docx not installed. Run: pip install python-docx"}
    except Exception as e:
        return {"error": str(e)}


# ── Tool 4: ocr_image ─────────────────────────────────────────────────────────
def ocr_image(file_path: str) -> dict:
    """
    Reads text from a scanned CV image using pytesseract (Tesseract OCR).
    Requires tesseract to be installed on the system.
    """
    try:
        import pytesseract
        from PIL import Image

        image = Image.open(file_path)

        # Improve OCR accuracy: convert to greyscale
        if image.mode != "L":
            image = image.convert("L")

        text = pytesseract.image_to_string(image, lang="eng")
        text = text.strip()

        if not text:
            return {
                "success": False,
                "error":   "OCR returned no text. Image may be too low resolution.",
            }

        return {
            "success":     True,
            "text":        text,
            "char_count":  len(text),
            "image_size":  image.size,
        }

    except ImportError:
        return {"error": "pytesseract or Pillow not installed. Run: pip install pytesseract pillow"}
    except Exception as e:
        return {"error": str(e)}


# ── Tool 5: detect_gaps ───────────────────────────────────────────────────────
def detect_gaps(experience: list) -> dict:
    """
    Finds periods longer than 3 months between jobs.
    Expects each item to have 'start' and 'end' as YYYY-MM strings.
    'end' can be 'present' for current role.
    """
    if not experience:
        return {"gaps": [], "note": "No experience provided"}

    def parse_date(s: str) -> datetime:
        s = s.strip().lower()
        if s in ("present", "current", "now", ""):
            return datetime.now().replace(day=1)
        # Try YYYY-MM
        try:
            return datetime.strptime(s, "%Y-%m")
        except ValueError:
            pass
        # Try YYYY
        try:
            return datetime.strptime(s, "%Y").replace(month=6)
        except ValueError:
            pass
        return None

    # Build list of (start, end, label) sorted by start date
    periods = []
    for job in experience:
        start = parse_date(job.get("start", ""))
        end   = parse_date(job.get("end",   "present"))
        if start and end:
            label = f"{job.get('role','?')} at {job.get('company','?')}"
            periods.append((start, end, label))

    if not periods:
        return {"gaps": [], "note": "Could not parse dates from experience"}

    periods.sort(key=lambda x: x[0])

    gaps = []
    for i in range(len(periods) - 1):
        end_of_this   = periods[i][1]
        start_of_next = periods[i + 1][0]

        if start_of_next > end_of_this:
            gap_months = (
                (start_of_next.year - end_of_this.year) * 12
                + (start_of_next.month - end_of_this.month)
            )
            if gap_months > 3:
                gaps.append({
                    "after_role":       periods[i][2],
                    "before_role":      periods[i + 1][2],
                    "gap_start":        end_of_this.strftime("%Y-%m"),
                    "gap_end":          start_of_next.strftime("%Y-%m"),
                    "gap_months":       gap_months,
                    "flag":             gap_months > 6,  # serious if > 6 months
                })

    return {
        "gaps":       gaps,
        "gap_count":  len(gaps),
        "has_gaps":   len(gaps) > 0,
    }


# ── Tool 6: detect_patterns ───────────────────────────────────────────────────
def detect_patterns(experience: list) -> dict:
    """
    Calculates average tenure across all roles.
    Flags job-hopping if average is under 18 months
    or if more than 2 roles were under 12 months.
    """
    if not experience:
        return {"avg_tenure_months": 0, "job_hopping": False, "note": "No experience"}

    def parse_date(s: str) -> datetime:
        s = s.strip().lower()
        if s in ("present", "current", "now", ""):
            return datetime.now().replace(day=1)
        try:
            return datetime.strptime(s, "%Y-%m")
        except ValueError:
            pass
        try:
            return datetime.strptime(s, "%Y").replace(month=6)
        except ValueError:
            pass
        return None

    tenures = []
    short_stints = []

    for job in experience:
        start = parse_date(job.get("start", ""))
        end   = parse_date(job.get("end", "present"))
        if start and end and end >= start:
            months = (end.year - start.year) * 12 + (end.month - start.month)
            label  = f"{job.get('role','?')} at {job.get('company','?')}"
            tenures.append(months)
            if months < 12:
                short_stints.append({"role": label, "tenure_months": months})

    if not tenures:
        return {"avg_tenure_months": 0, "job_hopping": False, "note": "Could not parse dates"}

    avg_tenure   = round(sum(tenures) / len(tenures), 1)
    job_hopping  = avg_tenure < 18 or len(short_stints) > 2

    return {
        "avg_tenure_months":  avg_tenure,
        "total_roles":        len(tenures),
        "short_stints":       short_stints,        # roles under 12 months
        "short_stint_count":  len(short_stints),
        "job_hopping":        job_hopping,
        "flag_reason": (
            f"Average tenure {avg_tenure} months (under 18)"
            if avg_tenure < 18 else
            f"{len(short_stints)} roles under 12 months"
            if len(short_stints) > 2 else
            "No job-hopping detected"
        ),
    }


# ── Tool 7: validate_output ───────────────────────────────────────────────────
def validate_output(cv_json: dict) -> dict:
    """
    Checks the final parsed CV JSON has all required fields.
    Returns a validation report so the LLM can fix anything missing.
    """
    required_top_keys = [
        "candidate", "skills", "experience",
        "education", "gaps", "patterns",
        "red_flags", "summary",
    ]
    required_candidate_keys = ["name", "email", "phone", "location"]

    issues = []

    # Check top-level keys
    for key in required_top_keys:
        if key not in cv_json:
            issues.append(f"Missing required field: '{key}'")

    # Check candidate sub-fields
    candidate = cv_json.get("candidate", {})
    if isinstance(candidate, dict):
        for key in required_candidate_keys:
            if not candidate.get(key):
                issues.append(f"candidate.{key} is empty or missing")
    else:
        issues.append("'candidate' must be an object")

    # Check lists are actually lists
    for list_field in ["skills", "experience", "education", "red_flags"]:
        val = cv_json.get(list_field)
        if val is not None and not isinstance(val, list):
            issues.append(f"'{list_field}' must be a list")

    # Check summary is non-empty
    summary = cv_json.get("summary", "")
    if not summary or len(summary) < 20:
        issues.append("'summary' is too short or empty")

    return {
        "valid":       len(issues) == 0,
        "issue_count": len(issues),
        "issues":      issues,
        "message":     "Output is valid — ready to return." if not issues else
                       f"{len(issues)} issue(s) found. Please fix before returning.",
    }